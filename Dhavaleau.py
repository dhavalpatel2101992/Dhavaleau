from flask import Flask, render_template,redirect,url_for,session,request,jsonify,json,g,send_file
from flask_session import Session
import pyodbc
from werkzeug.utils import secure_filename
import docx
import socket
import calendar
from time import strptime
import os
from datetime import timedelta
import pandas as pd
import time
import shutil
import pythoncom,win32com.client

# cnxn = pyodbc.connect("Driver={SQL Server};"
#                       "Server=CSTPTNSQLDEV1;"
#                       "Database=TST711CP;"
#                       "UID=cer_ro;"
#                       "PWD=C3Rro789!;")
# q_fisperiod = '''select distinct concat("FY_CD", '-',"PD_NO") as "Fiscal Period" from "DELTEK"."RPT_REVENUE_WKS" '''
# df_fisperiod = pd.read_sql_query(q_fisperiod, cnxn)
# df_fisperiod.to_csv('delete.csv')
# df_fisperiod = pd.read_csv('FiscalPeriod.csv')

orderdict = {'asc':1,'desc':0}
RW_FPAV = pd.read_csv('StaticParameters.csv')
RW_T_Project = pd.read_csv('projecttable.csv')
RW_T_Project['ProjectFilter']=RW_T_Project['PROJ_ID']+'.'+RW_T_Project['PROJ_NAME']
def RW_extractnumber(level):
    if level.startswith('Transaction'):
        return ''
    else:
        return level[6:]
RW_Parameters=['FiscalPeriod','ProjectLevel','Range','OrganizationLevel','Project']
RW_PDCV={'FiscalPeriod':'2020-02','ProjectLevel':'Level 3','Range':'Year To Date','OrganizationLevel':'Level 1','Project': None}
RW_PDAV = {'FiscalPeriod': RW_FPAV['FiscalPeriod'].tolist(),
                 'ProjectLevel':RW_FPAV['ProjectLevel'].dropna().tolist(),
                 'Range':RW_FPAV['Range'].dropna().tolist(),
                 'OrganizationLevel':RW_FPAV['OrganizationLevel'].dropna().tolist(),
                'Project':RW_T_Project[RW_T_Project['LVL_NO'].astype(str).str.contains(RW_extractnumber(RW_PDCV['ProjectLevel']))]['ProjectFilter'].tolist()}


def readtxt(filename):
    content = []
    counter = 0
    finaldflist = []
    doc = docx.Document(filename)
    for ix, para in enumerate(doc.paragraphs):
        if 'period' in para.text.lower():
            monthname = para.text.split(' ')[0].zfill(2)
            new = monthname[0].upper() + monthname[1:3].lower()
            monthnum = str(strptime(new, '%b').tm_mon)
            nextmonthnum = str(strptime(new, '%b').tm_mon + 1)
            year = para.text.split(' ')[1]
            nextyear = year
            print(monthname)
            print(new)
            print(nextmonthnum)
            print(year)

            if monthnum == '12':
                nextmonthnum = '1'
                nextyear = str(int(year) + 1)
            lastday = calendar.monthrange(int(year), int(monthnum))[1]
            vtable = doc.tables[counter]
            for vrows in vtable.rows:
                for vcell in vrows.cells:
                    cellcontent = []
                    for vpara in vcell.paragraphs:
                        for run in vpara.runs:
                            #                             print(run.text.strip()+'-'+str(run.font.color.rgb)+'-'+str(run.bold))
                            if 'FDM Period changed' in str(run.text.strip()):
                                content = []
                            if 'us holiday' not in run.text.lower():
                                cellcontent.append(run.text.strip())
                    cellcontent = list(filter(None, cellcontent))
                    content.append(list(dict.fromkeys(cellcontent)))
            #         print(content)
            counter -= 1
            df = pd.DataFrame(content)
            cols = list(df.columns)
            cols.remove(0)
            df['combined'] = df[cols].apply(lambda row: '&'.join(c for c in row.values.astype(str) if c != 'None'),
                                            axis=1)
            df['date'] = df[0].str.findall('\d\d')
            df['date'] = df['date'].apply(''.join)
            df['month'] = year + '-' + monthnum
            if max(df['date']) == str(lastday):
                df.loc[df.index[df['date'] == str(lastday)].values.tolist()[0] + 1:,
                'month'] = nextyear + '-' + nextmonthnum
            df['subject1'] = df[0].str.findall('\*.*\*')
            df['subject1'] = df['subject1'].apply(''.join)
            df['subject'] = df['subject1'] + df['combined']
            df['finaldatePST'] = df['month'] + '-' + df['date']
            #         print(df)
            df['finaldatePST'] = pd.to_datetime(df['finaldatePST'], format='%Y-%m-%d')
            df['finaldateIST'] = pd.to_datetime(df['finaldatePST'], format='%Y-%m-%d') + pd.Timedelta('1 day')
            df = df[df[1].str.len() > 0]
            finaldflist.append(df[['finaldateIST', 'subject']])
    finaldf=pd.concat(finaldflist)
    finaldf.columns = ['Date(IST)','Subject']
    return finaldf
def sendmeeting(row,RecepientID):  # 2019-01-16
    subject=row['Subject']
    # date=row['finaldateIST'].strftime('%Y-%m-%d')
    date=row['Date(IST)']
    pythoncom.CoInitialize()
    oOutlook = win32com.client.Dispatch("Outlook.Application")
    appt = oOutlook.CreateItem(1)  # 1 - olAppointmentItem
    appt.Start = date #+ ' 00:00'
    appt.Subject = subject
    appt.Body = '''\nCell Phone Contacts for GL Staff
    Amy Carniglia 619-665-0526 (Manager)
    Rosie Barrie 858-527-8883
    Fabio Braga 858-658-5448
    GAOC.GL@qti.qualcomm.com

    If we ever see issues with GL posts (e.g. flags not being set, multiple posts, unexpected results in the source) please use the following protocol:
    •  Contact the poster for that day.  The poster can be found on the GL monthly calendar that gets sent out by GL.
    •  If you do not get ahold of them within 30 minutes, please call manager’s cell phone.'''
    appt.Duration = 15
    appt.AllDayEvent = True
    # appt.Location = 'Office - Room 132A'
    appt.ReminderMinutesBeforeStart = 720
    appt.ReminderSet = True
    appt.MeetingStatus = 1
    appt.ResponseRequested = False
    appt.Recipients.Add(RecepientID)
    # Set Pattern, to recur every day, for the next 5 days
    # pattern = appt.GetRecurrencePattern()
    # pattern.RecurrenceType = 0
    # pattern.Occurrences = "5"
    appt.Save()
    appt.Send()
    print("Invite sent to "+ RecepientID+ " for "+date)


app = Flask(__name__)
app.debug = True
app.config['SECRET_KEY'] = os.urandom(24)
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes = 30)
app.config['SESSION_REFRESH_EACH_REQUEST']=True
sess = Session()
sess.init_app(app)

uploads_dir = os.path.join(app.instance_path, 'KSM_Temp_Files')
uploads_dir_glcalendar = os.path.join(app.instance_path, 'GL Calendar')
if not os.path.exists(uploads_dir):
    os.makedirs(uploads_dir)
if not os.path.exists(uploads_dir_glcalendar):
    os.makedirs(uploads_dir_glcalendar)


@app.before_request
def before_request():
    g.user=None
    if 'user' in session:
        g.user = session['user']


@app.route('/')
def landing():
    if g.user:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login/')
def login():
    if g.user:
        return redirect(url_for('dashboard'))
    return render_template("login.html")


@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if request.method=='POST':
        username = request.form.get('uname')
        # userpass = request.form.get('pass')
        # if (username == 'dhav' and userpass == 'pate'):
        if (username in ['dhavpate','sbairu','singlaa','c_vkmarri']):
            session['user'] =username
            session['RW_PCV'] = RW_PDCV.copy()
            session['RW_PAV'] = RW_PDAV.copy()
            session['parameter'] = {'nooffiles': 999 , 'successrate': 999, 'nooffailedfiles':999 }
            return redirect(url_for('obiee'))
        else:
            return redirect(url_for('login'))

    if ('user' in session and session['user'] == 'dhavpate'):
        return redirect(url_for('obiee'))
    return redirect(url_for('login'))

@app.route('/ksm', methods=["GET","POST"])
def ksm():
    if g.user:
        return render_template("ksm.html", parameter=session['parameter'] )
    return redirect(url_for('login'))

@app.route('/resetksm', methods=["GET","POST"])
def resetksm():
    session['parameter']['nooffiles'] = 999
    session['parameter']['nooffailedfiles'] = 999
    session['parameter']['successrate'] = 999
    return redirect(url_for('ksm'))

@app.route('/obiee')
def obiee():
    if g.user:
        return render_template("obiee.html")
    return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.pop('user')
    return redirect(url_for('login'))

@app.route('/obiee/RW',methods=['GET', 'POST'])
def RW():
    if g.user:
        return render_template("RW.html" , currentdict=session['RW_PCV'], valuedict = session['RW_PAV'])
    return redirect(url_for('login'))

@app.route('/obiee/RW/prompt' ,methods=['GET', 'POST'])
def RW_prompt():
    if request.method == 'POST':
        for pm in RW_Parameters:
            pv = request.form.get(pm)
            session['RW_PCV'][pm] = pv if len(pv) > 0 else RW_PDCV[pm]
        session['RW_PAV']['Project']=RW_T_Project[RW_T_Project['LVL_NO'].astype(str).str.contains(RW_extractnumber(session['RW_PCV']['ProjectLevel']))]['ProjectFilter'].tolist()

        query = 'select * from "DELTEK"."RPT_REVENUE_WKS" where FY_CD=\'{}\' and PD_NO=\'{}\' and PROJ_ID=\'{}\''.format(session['RW_PCV']['FiscalPeriod'][:4],int(session['RW_PCV']['FiscalPeriod'][5:]),session['RW_PCV']['Project'])
        # session['df'] = pd.read_sql_query(query, cnxn)
        print(query)
        session['df'] = pd.read_parquet('RPT_REVENUE_WKS.parq')


        return redirect(url_for('RW'))
@app.route('/obiee/RW/prompt_project', methods=['GET','POST'])
def RW_prompt_project():
    PrjLevel = request.args.get('ProjectLevel')
    retval=RW_T_Project[RW_T_Project['LVL_NO'].astype(str).str.contains(RW_extractnumber(PrjLevel))]['ProjectFilter'].tolist()
    return json.dumps(retval)


@app.route("/ksmupload", methods=["POST"])
def ksmupload():
    shutil.rmtree(uploads_dir)
    os.makedirs(uploads_dir)
    uploaded_files = request.files.getlist("file[]")
    session['parameter']['nooffiles']=len(uploaded_files)
    for file in uploaded_files:
        file.save(os.path.join(uploads_dir, file.filename))
    cmd="python validate.py "+uploads_dir+ " >"+os.path.join(uploads_dir,"log.txt")
    os.system(cmd)
    counter = len(uploaded_files)
    with open(os.path.join(uploads_dir,"log.txt"),'r') as f:
        for line in f:
            # word = line.split()
            if "File Complies with standards, no exceptions found!!!" in line:
                counter -= 1
    print(counter)
    session['parameter']['nooffailedfiles'] = counter
    session['parameter']['successrate'] = round(100*(1 - counter/session['parameter']['nooffiles']),2)
    return redirect(url_for('ksm'))

@app.route("/ksmdownloadvalidationlog")
def ksmdownloadvalidationlog():
    try:
        return send_file(os.path.join(uploads_dir,"log.txt"), attachment_filename='validation.txt', mimetype='text/plain',as_attachment=True,cache_timeout=0)
    except Exception as e:
        return str(e)

@app.route('/glposting')
def glposting():
    if g.user:
        return render_template("glposting.html",parameter=session['parameter'])
    session['preurl'] = request.url
    return redirect(url_for('login'))


@app.route('/qrsschedule')
def qrsschedule():
    if g.user:
        return render_template("QRSSchedule.html",parameter=session['parameter'])
    session['preurl'] = request.url
    return redirect(url_for('login'))

@app.route('/qrsschedule/filenames')
def qrs_files():
    period = request.args.get('period')
    dept = request.args.get('dept')

    def extractprocess(period, dept):
        files = []
        baseurl = 'https://spfin.qualcomm.com/sites/QRS'
        user = 'ap\dhavpate'
        password = 'canada@123'
        headers = {'accept': 'application/json;odata=verbose', 'Connection': 'close'}
        import requests
        from requests_ntlm import HttpNtlmAuth
        from bs4 import BeautifulSoup
        import urllib
        url = baseurl + '/' + period + '/' + dept
        #     print('Getting content from',url)
        resp = requests.get(url, auth=HttpNtlmAuth(user, password), headers=headers, verify=False)
        requests.session().close()
        if resp.status_code != requests.codes.ok:
            # raise Exception('Unexpected response from server: %s (%d).' % (resp.reason, resp.status_code))
            qrsfiles = ['File(s) Not Found ... Please do selection again !']
            session['qrsfiles'] = qrsfiles
            return qrsfiles
        soup = BeautifulSoup(resp.text, "html.parser")
        for link in soup.findAll('a'):
            href = link.get('href')
            title = link.string
            if title is not None and '?RootFolder=%2Fsites%2FQRS' + urllib.parse.quote('/' + period + '/' + dept + '/',
                                                                                       safe='') in href:
                files = files + extractprocess(period, dept + '/' + title)
            if title is not None and '.xlsx' in href:
                print(href)
                files.append(href)
        return files
    qrsfiles=extractprocess(period,dept)#['dhaval',period,dept]
    session['qrsfiles']=qrsfiles
    return json.dumps(qrsfiles)

@app.route('/qrsschedule/dataconsolidation', methods=['POST'])
def dataconsolidation():
    qrsfiles=session['qrsfiles']
    if qrsfiles[0] == 'File(s) Not Found ... Please do selection again !':
        return json.dumps({'Status': 'Process Failed ', 'files': qrsfiles})
    if len(qrsfiles)>0:
        print("Data Consolidation Started ...")
        from QRS_Schedule_Helper import dfgenerator,csvsavingfunc,excelsavingfunc,sendmailfunc
        df, processedfiles, failedfiles = dfgenerator(qrsfiles)
        if df is not None:
            csvsavingfunc(df, r'//canister/infa_corpfin_dev/edw9_r12/SrcFiles/EDW/QRS_ETL_SourceFile.csv')
            excelsavingfunc(df, 'QRS_File_For_Users.xlsx')
        sendmailfunc(processedfiles, failedfiles)
    return json.dumps({'Status':'Process Completed ... Email Sent !','files':qrsfiles})

@app.route('/glpostingfilesaving', methods=['POST'])
def glpostingfilesaving():
    shutil.rmtree(uploads_dir_glcalendar)
    os.makedirs(uploads_dir_glcalendar)
    uploaded_docx = request.files.get("inputfile")
    path=os.path.join(uploads_dir_glcalendar, uploaded_docx.filename)
    session['parameter']['glcalendarfilename'] = uploaded_docx.filename
    session['parameter']['glcalendarpath'] = path
    uploaded_docx.save(path)
    session['parameter']['glcalendartable'] = readtxt(session['parameter']['glcalendarpath'])
    return redirect(url_for('glposting'))

@app.route('/glpostingsendinvitation', methods=['POST'])
def glpostingsendinvitation():
    tbl = pd.DataFrame(request.json[0])
    lst= request.json[1]
    if len(lst)==3:
        print(lst[0]['value'])
        print(lst[1]['value'])
        tbl.apply(sendmeeting,args=[lst[1]['value']],axis=1)
    return json.dumps({'Status':'Process Completed Successfully ...'})

if __name__ == "__main__":
    hostname = socket.gethostname()
    IPAddr = socket.gethostbyname(hostname)
    app.run(threaded=True, host=IPAddr)
    # app.run(threaded=True)

